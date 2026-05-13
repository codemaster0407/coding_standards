import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_json(json_path, output_path):
    # Load the JSON data
    with open(json_path, 'r') as f:
        data = json.load(f)
    
    # Create a new Document
    doc = Document()
    
    # Optional: Set a main title if the JSON doesn't start with one
    # doc.add_heading('Coding Standards Document', 0)

    for section in data.get('sections', []):
        header_type = section.get('header_type', 'p')
        header_text = section.get('header_text', '')
        text_content = section.get('text', '')

        # Map header types to docx levels and set sizes
        if header_type == 'h1':
            h = doc.add_heading(header_text, level=1)
            for run in h.runs:
                run.font.size = Pt(11)
        elif header_type == 'h2':
            h = doc.add_heading(header_text, level=2)
            for run in h.runs:
                run.font.size = Pt(11)
        elif header_type == 'h3':
            h = doc.add_heading(header_text, level=3)
            for run in h.runs:
                run.font.size = Pt(11)
        elif header_type == 'h4':
            h = doc.add_heading(header_text, level=4)
            for run in h.runs:
                run.font.size = Pt(11)
        else:
            # For p or others, just add a bold line or similar if header exists
            if header_text:
                p = doc.add_paragraph()
                p.add_run(header_text).bold = True
        
        # Add the main text content
        if text_content:
            # Split text by newline first
            raw_lines = text_content.split('\n')
            for raw_line in raw_lines:
                clean_raw_line = raw_line.strip()
                if not clean_raw_line:
                    continue
                
                # If the line contains multiple bullet markers (e.g. "* Item 1 * Item 2")
                # we split them. Using regex to find " * " or starting "*"
                import re
                # Split by '*' but keep the content. We look for '*' at start or space+*
                bullet_parts = re.split(r'(?<=^)\*\s*|(?<=\s)\*\s*', clean_raw_line)
                
                # Remove empty parts
                bullet_parts = [p.strip() for p in bullet_parts if p.strip()]
                
                if clean_raw_line.startswith(('*', '-', '•')) or len(bullet_parts) > 1:
                    for part in bullet_parts:
                        p = doc.add_paragraph(part, style='List Bullet')
                        for run in p.runs:
                            run.font.size = Pt(11)
                else:
                    # Add as a normal paragraph
                    p = doc.add_paragraph(clean_raw_line)
                    for run in p.runs:
                        run.font.size = Pt(10)

    # Save the document
    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

