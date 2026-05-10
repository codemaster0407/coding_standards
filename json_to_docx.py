import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_json(json_path, output_path):
    # Load the JSON data
    with open(json_path, 'r') as f:
        data = json.load(f)
    
    # Create a new Document
    doc = Document()
    
    # Optional: Set a main title if the JSON doesn't start with one
    # doc.add_heading('Coding Standards Document', 0)

    for section in data.get('sections', []):
        header_type = section.get('header_type', 'p')
        header_text = section.get('header_text', '')
        text_content = section.get('text', '')

        # Map header types to docx levels
        if header_type == 'h1':
            doc.add_heading(header_text, level=1)
        elif header_type == 'h2':
            doc.add_heading(header_text, level=2)
        elif header_type == 'h3':
            doc.add_heading(header_text, level=3)
        elif header_type == 'h4':
            doc.add_heading(header_text, level=4)
        else:
            # For p or others, just add a bold line or similar if header exists
            if header_text:
                p = doc.add_paragraph()
                p.add_run(header_text).bold = True
        
        # Add the main text content
        if text_content:
            doc.add_paragraph(text_content)

    # Save the document
    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == "__main__":
    json_file = 'coding_standards_extracted.json'
    output_docx = 'Coding_Standards.docx'
    create_docx_from_json(json_file, output_docx)
